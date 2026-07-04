"""Unit tests for ParsingService — tests real parsing logic on real bytes."""
import io
import pytest
from app.services.parsing_service import ParsingService
from app.core.exceptions import AppException


@pytest.fixture
def service():
    return ParsingService()


def test_parse_txt_returns_content(service):
    content = b"Hello world. This is a test document."
    pages = service.parse("txt", content=content)
    assert len(pages) == 1
    assert "Hello world" in pages[0].text


def test_parse_txt_raises_on_empty(service):
    with pytest.raises(AppException):
        service.parse("txt", content=b"   ")


def test_parse_markdown_extracts_sections(service):
    md = b"# Section One\nSome content here.\n## Section Two\nMore content."
    pages = service.parse("markdown", content=md)
    sections = [p.section for p in pages]
    assert "Section One" in sections
    assert "Section Two" in sections


def test_parse_markdown_raises_on_empty(service):
    with pytest.raises(AppException):
        service.parse("markdown", content=b"")


def test_parse_pdf_extracts_text(service):
    """Create a minimal valid PDF in memory and parse it."""
    try:
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 750, "Test PDF content for parsing.")
        c.save()
        buf.seek(0)
        pages = service.parse("pdf", content=buf.read())
        assert len(pages) >= 1
        assert any("Test PDF" in p.text for p in pages)
    except ImportError:
        pytest.skip("reportlab not installed — skipping PDF generation test")


def test_parse_docx_extracts_text(service):
    """Build a minimal DOCX in memory and parse it."""
    from docx import Document as DocxDoc
    buf = io.BytesIO()
    doc = DocxDoc()
    doc.add_paragraph("This is a test paragraph in a DOCX file.")
    doc.add_heading("Section Heading", level=1)
    doc.add_paragraph("Content after the heading.")
    doc.save(buf)
    buf.seek(0)
    pages = service.parse("docx", content=buf.read())
    full_text = " ".join(p.text for p in pages)
    assert "test paragraph" in full_text


def test_parse_docx_extracts_section_headings(service):
    from docx import Document as DocxDoc
    buf = io.BytesIO()
    doc = DocxDoc()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("Content of the introduction section.")
    doc.save(buf)
    buf.seek(0)
    pages = service.parse("docx", content=buf.read())
    sections = [p.section for p in pages if p.section]
    assert "Introduction" in sections


def test_unsupported_source_type_raises(service):
    with pytest.raises(AppException):
        service.parse("xlsx", content=b"data")
